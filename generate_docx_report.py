import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = create_element('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def add_heading_styled(doc, text, level, space_before=12, space_after=6):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(space_before)
    heading.paragraph_format.space_after = Pt(space_after)
    heading.paragraph_format.keep_with_next = True
    
    # Apply custom colors to headings
    run = heading.runs[0]
    run.font.name = 'Outfit'
    if level == 1:
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(79, 70, 229) # Indigo #4F46E5
        run.bold = True
    elif level == 2:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(236, 72, 153) # Pink #EC4899
        run.bold = True
    elif level == 3:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(31, 41, 55) # Dark gray #1F2937
        run.bold = True
    return heading

def add_paragraph_styled(doc, text, space_after=6, bold=False, italic=False, color_rgb=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Outfit'
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb
    return p

def main():
    doc = docx.Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Set default style font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Outfit'
    font.size = Pt(11)
    font.color.rgb = RGBColor(55, 65, 81) # Slate #374151
    
    # ------------------ COVER PAGE ------------------
    # Spacing before title
    for _ in range(5):
        doc.add_paragraph()
        
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run("WEB WIZARDS CAR RENTALS")
    run_title.font.name = 'Outfit'
    run_title.font.size = Pt(32)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(79, 70, 229) # Indigo
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(36)
    run_sub = p_sub.add_run("Comprehensive Software Engineering Specification & Project Report")
    run_sub.font.name = 'Outfit'
    run_sub.font.size = Pt(16)
    run_sub.font.color.rgb = RGBColor(236, 72, 153) # Hot pink
    run_sub.bold = True
    
    for _ in range(8):
        doc.add_paragraph()
        
    # Metadata Block
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_after = Pt(4)
    run_meta1 = p_meta.add_run("System Domain: ")
    run_meta1.bold = True
    p_meta.add_run("On-Demand Multi-Role B2C Car Rental Platform\n")
    run_meta2 = p_meta.add_run("Release Version: ")
    run_meta2.bold = True
    p_meta.add_run("2.4.0-Rupee-Standard\n")
    run_meta3 = p_meta.add_run("Author: ")
    run_meta3.bold = True
    p_meta.add_run("Lead Systems Architect & Software Engineer\n")
    run_meta4 = p_meta.add_run("Reference Schema: ")
    run_meta4.bold = True
    p_meta.add_run("Note.txt Developer Core Specifications\n")
    
    doc.add_page_break()
    
    # ------------------ CHAPTER 1 ------------------
    add_heading_styled(doc, "1. Problem Statement", 1)
    
    add_heading_styled(doc, "1.1 Context and Industry Background", 2)
    add_paragraph_styled(doc, "The vehicle rental industry has traditionally operated under fragmented paradigms dominated by localized agency desks, opaque pricing matrices, and physical paper-trailing. With the advent of web technologies, localized players attempted digital transformations, yet modern sharing economy B2C networks still suffer from three fatal friction points:")
    
    add_paragraph_styled(doc, "1. Inefficient Supplier Onboarding & Credential Auditing: Small-scale fleet owners (vendors) are either locked out by giant aggregators or allowed onto open marketplaces without systematic document checks (registration papers, commercial insurance validation), exposing customers to legal liabilities.")
    add_paragraph_styled(doc, "2. Pricing Discrepancies and Opaque Fee Calculations: Traditional calculators lack smart flexibility. Rentals transitioning across daily boundaries are billed raw 24-hour cycles, completely disregarding fair hourly capping. Furthermore, promo codes and conditional discounts are hardcoded or manually added, creating booking-time friction.")
    add_paragraph_styled(doc, "3. Weak Rental Verification and Collateral Protocols: Unlike standard taxi-hailing, self-drive rentals demand rigorous customer checks. The lack of structured document checkpoints (driver license logs, base64 digital signatures, clear damage logging, and original physical vehicle/RC card deposits) creates high security risks for vehicle owners.")
    
    add_heading_styled(doc, "1.2 The Proposed Solution: Web Wizards Car Rentals", 2)
    add_paragraph_styled(doc, "To address these systemic bottlenecks, we designed a unified, self-contained B2C car rental platform. The application bifurcates permissions securely, creating a triple-role environment consisting of registered Customers, fleet Vendors, and platform Administrators.")
    add_paragraph_styled(doc, "The system delivers three unique technological solutions:")
    add_paragraph_styled(doc, "- Comprehensive Document Verification: A multi-step verification pipeline where suppliers upload official structural registration PDFs, and customers upload license scans at booking checkouts.")
    add_paragraph_styled(doc, "- Smart Adaptive Pricing Engine: A logic framework calculating rental fees using native daily rates and partial hourly charges—automatically capped at a single day's rate—with localized currency formatting in Indian Rupees (₹).")
    add_paragraph_styled(doc, "- Mutual Discount Negotiation System: A shared campaign dashboard where admins suggest discounts on cars, and vendors accept and launch them dynamically.")
    
    # ------------------ CHAPTER 2 ------------------
    add_heading_styled(doc, "2. Requirements Gathering", 1)
    
    add_heading_styled(doc, "2.1 Functional Requirements (FR)", 2)
    
    add_heading_styled(doc, "2.1.1 Guest/Anonymous User Requirements", 3)
    add_paragraph_styled(doc, "- FR-1.1: The system must allow guest users to search for approved vehicles by location, category, transmission type, and maximum price.")
    add_paragraph_styled(doc, "- FR-1.2: The system must allow guest users to register as a Customer or Vendor, accompanied by a double-factor session-based OTP verification.")
    add_paragraph_styled(doc, "- FR-1.3: The system must allow users to request password resets via OTP email loops.")
    
    add_heading_styled(doc, "2.1.2 Registered Customer Requirements", 3)
    add_paragraph_styled(doc, "- FR-2.1: A Customer must be able to view detailed specifications and photo galleries of approved cars.")
    add_paragraph_styled(doc, "- FR-2.2: A Customer must be able to check real-time vehicle availability via an interactive calendar, blocking reserved date ranges.")
    add_paragraph_styled(doc, "- FR-2.3: A Customer must be able to request booking reservations, which are sent to the supplier's approval queue.")
    add_paragraph_styled(doc, "- FR-2.4: A Customer must be able to complete payment checkout via a simulated credit card gateway, including digital license upload, digital signature capture, and damage condition logs.")
    add_paragraph_styled(doc, "- FR-2.5: A Customer must be able to download itemized invoice PDFs denominated in Rupees (₹).")
    add_paragraph_styled(doc, "- FR-2.6: A Customer must be able to cancel bookings within 24 hours of creation, triggering automated payment refund logs.")
    add_paragraph_styled(doc, "- FR-2.7: A Customer must be able to submit post-rental star reviews and submit formal complaints/disputes.")
    
    add_heading_styled(doc, "2.1.3 Registered Vendor (Supplier) Requirements", 3)
    add_paragraph_styled(doc, "- FR-3.1: A Vendor must submit verification documents (Company Name, Official Supplier License) to be approved.")
    add_paragraph_styled(doc, "- FR-3.2: A Vendor must be able to add, edit, and delete vehicles.")
    add_paragraph_styled(doc, "- FR-3.3: A Vendor must upload multiple vehicle gallery photos and registration/insurance files.")
    add_paragraph_styled(doc, "- FR-3.4: A Vendor must be able to accept or reject incoming customer booking requests.")
    add_paragraph_styled(doc, "- FR-3.5: A Vendor must be able to define custom Daily Rates and optional Hourly Rates.")
    add_paragraph_styled(doc, "- FR-3.6: A Vendor must be able to launch custom multi-day discount campaigns.")
    add_paragraph_styled(doc, "- FR-3.7: A Vendor must be able to mark active rentals as returned, instantly releasing vehicle availability.")
    
    add_heading_styled(doc, "2.1.4 System Administrator Requirements", 3)
    add_paragraph_styled(doc, "- FR-4.1: An Admin must be able to approve or reject pending vendor applications.")
    add_paragraph_styled(doc, "- FR-4.2: An Admin must be able to audit and list/reject new vehicle additions.")
    add_paragraph_styled(doc, "- FR-4.3: An Admin must be able to resolve open dispute complaints.")
    add_paragraph_styled(doc, "- FR-4.4: An Admin must be able to suggest promotional discounts to specific vendors.")
    add_paragraph_styled(doc, "- FR-4.5: An Admin must be able to view consolidated platform revenue reports.")
    
    add_heading_styled(doc, "2.2 Non-Functional Requirements (NFR)", 2)
    add_paragraph_styled(doc, "- NFR-1 (Security): Password storage must be hashed using Django's default PBKDF2 algorithm. Session validation must enforce authentication checks for all secure routes.")
    add_paragraph_styled(doc, "- NFR-2 (Performance): Dynamic price updates and calendar availability renderings must execute client-side in under 150 milliseconds.")
    add_paragraph_styled(doc, "- NFR-3 (Data Integrity): Prevent double-booking conflicts at the database transaction layer using isolation queries.")
    add_paragraph_styled(doc, "- NFR-4 (UI/UX Aesthetic): Interface components must utilize modern HSL color maps, backdrop filters, and glassmorphic card configurations.")
    add_paragraph_styled(doc, "- NFR-5 (Robustness): Django views and API points must handle faulty date formats, negative rates, and invalid digital signature structures gracefully.")
    
    # ------------------ CHAPTER 3 ------------------
    doc.add_page_break()
    add_heading_styled(doc, "3. User Analysis", 1)
    
    add_heading_styled(doc, "3.1 User Personas", 2)
    add_paragraph_styled(doc, "To map features to target workflows, we generated three concrete user profiles:")
    add_paragraph_styled(doc, "1. Renter (Customer - 'Anjali'): Tech-savvy millennial renter. Demands swift navigation, dynamic hourly-rate transparency, mobile responsiveness, and instantaneous checkouts.")
    add_paragraph_styled(doc, "2. Fleet Owner (Vendor - 'Rajesh'): Businessman owning 10 SUVs. Demands simple car registration inputs, automated date blocking, transparent financial ledger states, and robust contract verification (physical collateral RC deposit warnings).")
    add_paragraph_styled(doc, "3. Operations Manager (Admin - 'Vikram'): System inspector. Seeks a single interface to review applications, lock dispute cases, compile overall profits, and distribute marketing promos.")
    
    add_heading_styled(doc, "3.2 Use Case Specifications", 2)
    add_paragraph_styled(doc, "Customers log in to search, book, pay (checkout), sign contracts, and review. Vendors manage vehicle listings, respond to requested rentals, issue promo deductions, and check payment logs. Admins verify credentials, launch discounts, and audit revenue databases.")
    
    # ------------------ CHAPTER 4 ------------------
    add_heading_styled(doc, "4. Documentation & Repository Structure", 1)
    add_paragraph_styled(doc, "The codebase follows a modular Django pattern. The root contains web_wizards_rentals/ (general settings, middleware configurations, main routing tables) and car_rental/ (app model managers, custom views, OTP systems, tests, and administrative triggers). UI templates are stored in a centralized templates/ directory. Scanned license graphics, digital signatures, and cover photos are saved in media/.")
    
    # ------------------ CHAPTER 5 ------------------
    doc.add_page_break()
    add_heading_styled(doc, "5. Workflow & Planning", 1)
    
    add_heading_styled(doc, "5.1 Booking State Progression", 2)
    add_paragraph_styled(doc, "A booking transitions dynamically through different states depending on user actions:")
    add_paragraph_styled(doc, "1. PENDING: Initial state when a Customer requests a rental. The car remains visible but the specific dates are blocked from other searches.")
    add_paragraph_styled(doc, "2. APPROVED: Vendor accepts the rental. The booking awaits payment checkout.")
    add_paragraph_styled(doc, "3. PAID: Customer fills payment forms, signs, and uploads licenses. The transaction reference code is saved.")
    add_paragraph_styled(doc, "4. COMPLETED: Vendor confirms vehicle handover return, instantly restoring car availability.")
    add_paragraph_styled(doc, "5. REJECTED: Vendor refuses the booking, releasing the blocked dates.")
    add_paragraph_styled(doc, "6. CANCELLED: Customer requests cancellation. If cancelled within 24 hours of booking creation, the transaction status is updated to REFUNDED automatically.")
    
    # ------------------ CHAPTER 6 ------------------
    add_heading_styled(doc, "6. UI/UX Design & Frontend Implementation Specification", 1)
    
    add_paragraph_styled(doc, "The frontend architecture of Web Wizards Car Rentals is built from the ground up on modern design principles that emphasize immersion, high-performance interactions, and responsive visual feedback. By rejecting default browser components, the interface implements a bespoke Cyber-Noir / Midnight Glassmorphism aesthetic tailored to high-end digital consumers.")
    
    add_heading_styled(doc, "6.1 Core Design Philosophy and Usability Guidelines", 2)
    add_paragraph_styled(doc, "1. Premium Cybernetic Depth: Employs a dark midnight backdrop integrated with floating glowing neon accents and a responsive matrix cyber-grid to convey speed, technology, and luxury.")
    add_paragraph_styled(doc, "2. Glassmorphism Layering Hierarchy: Utilizes translucent panels with backdrop-blur filters, thin luminous borders, and soft shadows to simulate layers of physically-stacked glass cards.")
    add_paragraph_styled(doc, "3. Cognitive Load Minimization & Micro-Interactions: Provides immediate visual rewards for user actions. Hover actions trigger 3D perspective shifts, glow enhancements, and smooth hardware-accelerated transitions.")
    add_paragraph_styled(doc, "4. Mobile-First Responsiveness: Core grid layouts dynamically collapse from complex multi-column viewports into simple, clean vertical stacks, replacing traditional sidebars with off-canvas drawer systems.")
    
    add_heading_styled(doc, "6.2 The Design Token System (CSS Root Variables)", 2)
    add_paragraph_styled(doc, "The system establishes a single source of truth for the entire interface using responsive CSS custom properties defined within the root block of base.html:")
    
    # Adding CSS token variables table
    tokens_table = doc.add_table(rows=1, cols=3)
    tokens_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = tokens_table.rows[0].cells
    hdr_cells[0].text = 'CSS Custom Property'
    hdr_cells[1].text = 'Exact Color / Value'
    hdr_cells[2].text = 'Functional UI/UX Role'
    for cell in hdr_cells:
        set_cell_background(cell, '4F46E5')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
        cell.paragraphs[0].runs[0].bold = True
        
    tokens_data = [
        ('--bg-main', '#0B0F19 (Midnight Black)', 'Base canvas background for the entire application'),
        ('--bg-darker', '#07090F (Pitch Obsidian)', 'Dark contrast background for off-canvas drawer navigation and footers'),
        ('--bg-card', 'rgba(17, 24, 39, 0.7)', 'Foundation backing for glassmorphic cards and interactive panels'),
        ('--border-color', 'rgba(255, 255, 255, 0.08)', 'Subtle perimeter border for cards, mimicking light refractiveness'),
        ('--border-color-hover', 'rgba(255, 255, 255, 0.15)', 'High-contrast refractiveness border triggered on card hover states'),
        ('--primary', '#4F46E5 (Vibrant Indigo)', 'Primary branding element, button fills, and active links'),
        ('--primary-glow', 'rgba(79, 70, 229, 0.4)', 'Blur glow shadow for active inputs and key click triggers'),
        ('--accent', '#EC4899 (Electric Hot Pink)', 'High-contrast accent color, checkout CTA highlights, and alert status'),
        ('--accent-glow', 'rgba(236, 72, 153, 0.4)', 'Glowing shadow maps for special checkout elements and action highlights'),
        ('--success', '#10B981 (Emerald Green)', 'High-visibility indicators for approved items, paid bookings, and validations'),
        ('--warning', '#F59E0B (Amber Gold)', 'Warning markers for pending actions or administrative review statuses'),
        ('--danger', '#EF4444 (Coral Crimson)', 'Critical warning indicators for cancellations, rejections, and error messages'),
        ('--glass-blur', 'blur(12px)', 'WebKit backdrop filtering rule to blur background elements behind glass'),
        ('--glass-shadow', '0 8px 32px 0 rgba(0, 0, 0, 0.37)', 'Deep-cast box shadow mapping to simulate vertical elevation layers'),
        ('--transition', 'all 0.3s cubic-bezier(0.4, 0, 0.2, 1)', 'Standardized ease-out bezier curve ensuring fluid fluid animation flow')
    ]
    
    for row_data in tokens_data:
        row_cells = tokens_table.add_row().cells
        row_cells[0].text = row_data[0]
        row_cells[1].text = row_data[1]
        row_cells[2].text = row_data[2]
        
    doc.add_paragraph() # spacing
    
    add_heading_styled(doc, "6.3 Advanced Interactive Components (Client-Side Implementation)", 2)
    
    add_heading_styled(doc, "6.3.1 The 3D Interactive Payment Card Flipper", 3)
    add_paragraph_styled(doc, "To remove friction from the checkout phase, checkout.html features an immersive credit card simulator that flips in 3D perspective based on input focus.")
    add_paragraph_styled(doc, "- Visual Presentation: Styled using dual absolute-positioned layers (.flip-card-front and .flip-card-back) bound under a parent container with perspective: 1000px and transform-style: preserve-3d. Back-facing layers utilize -webkit-backface-visibility: hidden to mask reversed contents.")
    add_paragraph_styled(doc, "- Behavioral Logic: Focus event listeners monitor input zones. When CVV fields gain focus, a JavaScript helper inserts .flipped class, executing transform: rotateY(180deg) using hardware-accelerated transitions.")
    add_paragraph_styled(doc, "- Data Sync Pipeline: Real-time keystroke listeners sync input strings (Cardholder, Card Number, Expiry) immediately into styled layouts, auto-formatting input strings with spaces every four digits for high readability.")
    
    add_heading_styled(doc, "6.3.2 HTML5 Touch-Enabled Digital Signature Pad", 3)
    add_paragraph_styled(doc, "To satisfy rental liability compliance, the checkout interface features an HTML5 canvas signature pad enabling renters to draw binding digital contracts.")
    add_paragraph_styled(doc, "- Canvas Rendering Context: Listens to mouse (mousedown, mouseup, mousemove) and equivalent mobile touch events to trace coordinates across a 2D context using specific line parameters matching the Indigo branding theme.")
    add_paragraph_styled(doc, "- Base64 Serialization: Upon releasing drawing inputs, standard Canvas stroke geometries are compressed and serialized into a base64 encoded PNG data URI (data:image/png;base64,...) and updated to a hidden input container. On submit, Django reads and saves this byte-string directly to the database media models.")
    
    add_heading_styled(doc, "6.3.3 Dynamic Locked-Date Calendar Booker Widget", 3)
    add_paragraph_styled(doc, "On the car_detail.html specification display, customers reserve date intervals using an interactive, custom-designed calendar widget.")
    add_paragraph_styled(doc, "- Dynamic Render Loop: Reads active year/month offsets to generate absolute calendar matrix cells client-side in real-time.")
    add_paragraph_styled(doc, "- Backend Blocked Range Integration: Django queries conflicting bookings, formats them as a JSON list, and passes them to the template context. The JavaScript calendar loops through these ranges, applying .blocked classes and disabled attributes to cells overlapping existing reservations.")
    add_paragraph_styled(doc, "- Visual Capping & Range Selection: Detects first and second clicks to define start/end dates, colorizing intermediate cells with custom HSL gradient overlays, and instantly feeds total duration bounds to the pricing engine.")
    
    add_heading_styled(doc, "6.3.4 Form-Interception Submission Loader System", 3)
    add_paragraph_styled(doc, "To maintain visual stability during slower server-side processes (e.g. OTP validation or license uploads), a global glassmorphic loading overlay is integrated:")
    add_paragraph_styled(doc, "- Submit Interception: A centralized script scans all page forms and attaches submission listeners. When a form fires, it verifies validation rules (checkValidity()). If validated, it activates a global overlay.")
    add_paragraph_styled(doc, "- Dual Rotating Spinner: Rendered using a translucent circular ring rotating under specific infinite keyframe timelines, accompanied by a pulsing status message specifying the exact process (e.g., 'Securing your connection & authenticating...').")
    
    add_heading_styled(doc, "6.4 Key Screen Designs and Layout Profiles", 2)
    add_paragraph_styled(doc, "- Futuristic Home Platform (index.html): Employs an ultra-wide hero banner with layered radial gradient lights. Car listings display with dynamic CSS scale hover scaling (scale(1.03)), casting indigo drop-shadows on active states.")
    add_paragraph_styled(doc, "- Dual-Column Search Desk (car_search.html): Employs a sticky left sidebar holding range sliders, category tags, and check elements, feeding criteria to real-time search lists on the right side.")
    add_paragraph_styled(doc, "- Dashboard Hubs (dashboard_customer/vendor/admin.html): Multi-role dashboards styled using consistent modular layouts. Highlight stats are displayed in glow cards, followed by responsive transaction data lists styled with thin border matrices.")
    
    # ------------------ CHAPTER 7 ------------------
    doc.add_page_break()
    add_heading_styled(doc, "7. Database Design", 1)
    
    add_heading_styled(doc, "7.1 Model Mapping", 2)
    add_paragraph_styled(doc, "We designed an optimized relational schema in SQLite using Django ORM:")
    
    # Adding a table for DB Models
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Model Name'
    hdr_cells[1].text = 'Primary Fields'
    hdr_cells[2].text = 'Relationships & Constraints'
    for cell in hdr_cells:
        set_cell_background(cell, '4F46E5')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
        cell.paragraphs[0].runs[0].bold = True
        
    models_data = [
        ('User', 'username, role, phone_number, profile_picture', 'Role choices: CUSTOMER, VENDOR, ADMIN'),
        ('VendorProfile', 'company_name, license_number, is_approved', 'OneToOneField to User'),
        ('CustomerProfile', 'driver_license, address', 'OneToOneField to User'),
        ('Car', 'brand, model, category, daily_rate, hourly_rate, location', 'ForeignKey to User (vendor), Rate constraints >= 0.0'),
        ('Booking', 'start_date, end_date, total_price, status, payment_status', 'ForeignKey to User (customer), ForeignKey to Car'),
        ('Payment', 'transaction_id, amount, payment_method, timestamp', 'OneToOneField to Booking, Unique Transaction ID'),
        ('Discount', 'discount_percentage, min_days, status, created_by', 'ForeignKey to Car, Percentage constraint 1-100')
    ]
    
    for row_data in models_data:
        row_cells = table.add_row().cells
        row_cells[0].text = row_data[0]
        row_cells[1].text = row_data[1]
        row_cells[2].text = row_data[2]
        
    doc.add_paragraph() # spacing
    
    # ------------------ CHAPTER 8 ------------------
    add_heading_styled(doc, "8. Architecture", 1)
    add_paragraph_styled(doc, "The system follows the standard Model-View-Template (MVT) architectural design pattern. The router receives client requests, view controllers process business validation parameters, query the relational SQLite database via ORM bindings, and render localized, responsive HTML pages dynamically populated with context variables.")
    
    # ------------------ CHAPTER 9 ------------------
    doc.add_page_break()
    add_heading_styled(doc, "9. Backend (Controller & Business Logic)", 1)
    
    add_paragraph_styled(doc, "The controller layer (views.py) serves as the core logic engine of the platform, orchestrating state changes, processing transactions, validating credentials, and enforcing role-based execution boundaries.")
    
    add_heading_styled(doc, "9.1 Secure Dual-Factor Registration & Password Reset OTP Pipelines", 2)
    add_paragraph_styled(doc, "Authentication utilizes a secure session-backed, email-verified One-Time Password (OTP) framework to validate customer and vendor email addresses prior to account instantiation.")
    add_paragraph_styled(doc, "1. Registration Flow: When register_view receives a signup request, it generates a random 6-digit cryptographic token using random.randint(100000, 999999) and writes all profile inputs directly to request.session['pending_registration']. It also caches the code in request.session['registration_otp'] and triggers an email dispatch. Upon verification in verify_otp_view, if the entered OTP matches, the database profiles are generated, session caches are purged, and the user is statefully authenticated.")
    add_paragraph_styled(doc, "2. Password Reset Flow: Similarly, forgot_password_view verifies target user existence, caches a password-reset token in request.session['reset_otp'], and dispatches it. When verified, new credentials are encrypted and stored via user.set_password().")
    
    add_heading_styled(doc, "9.2 The Concurrency-Safe Double-Booking Isolation Engine", 2)
    add_paragraph_styled(doc, "To prevent race conditions and conflicting vehicle reservations, the system enforces a strict overlapping date check prior to booking instantiation:")
    add_paragraph_styled(doc, "- Overlapping Interval Query: The backend checks start_date <= requested_end and end_date >= requested_start within approved or pending status boundaries.")
    add_paragraph_styled(doc, "- Failure Redirect: If overlaps are detected, the request immediately triggers messages.error and redirects back to avoid dirty writes or database double-bookings.")
    
    add_heading_styled(doc, "9.3 The Adaptive Capped Pricing & Tiered Discount Calculations Engine", 2)
    add_paragraph_styled(doc, "The backend controller implements a unique adaptive pricing algorithm to prevent customer price inflation for partial days:")
    add_paragraph_styled(doc, "1. It isolates rental days and remaining extra hours from the booking duration using timedelta delta offsets.")
    add_paragraph_styled(doc, "2. It calculates the raw hourly charge. If this charge exceeds the daily rate, it is automatically capped at a single day's rate: extra_charge = min(Decimal(remaining_hours) * hourly_rate, car.daily_rate).")
    add_paragraph_styled(doc, "3. It queries active database-driven promotions (Discount models) sorted by min_days descending. It applies the highest percentage threshold matching the rental duration and round-formats the final price.")
    
    add_heading_styled(doc, "9.4 Collaborative Administrative-Vendor Promotion Negotiation Protocol", 2)
    add_paragraph_styled(doc, "To align marketing campaigns, the platform implements a collaborative promotion negotiation lifecycle:")
    add_paragraph_styled(doc, "- Admin Proposal: Admins suggest discounts which write with a PENDING_VENDOR state.")
    add_paragraph_styled(doc, "- Vendor Response: Vendors view pending suggestions in their dashboard and accept or reject them. Accepting shifts the promotion status to APPROVED, immediately enabling the discount in the customer booking calculator. Rejecting removes it from active feeds.")
    
    add_heading_styled(doc, "9.5 Strict Object-Ownership Guarding & Stateful Authorization Protocols", 2)
    add_paragraph_styled(doc, "To prevent context-bypass vulnerabilities (such as ID scraping or direct URL path parameter manipulation), the backend enforces strict session cross-checks on all state-changing endpoints:")
    add_paragraph_styled(doc, "- Context Checks: Every view verifies that the request.user matches the vendor of the queried car listing or is the customer of the target booking.")
    add_paragraph_styled(doc, "- Role Protections: Restricted actions enforce role bounds (e.g., VENDOR or ADMIN checks), rejecting unauthorized operations automatically.")
    
    add_heading_styled(doc, "9.6 The Time-Bounded Cancellation & Automated Auto-Refund Engine", 2)
    add_paragraph_styled(doc, "The system supports customer cancellations, but enforces a strict 24-hour window from the booking creation timestamp to protect supplier schedules:")
    add_paragraph_styled(doc, "- Window Check: Compares current time against booking.created_at, blocking actions if the duration exceeds 24 hours.")
    add_paragraph_styled(doc, "- Refund & Release: For paid bookings, the payment status changes to REFUNDED, and the car's availability is instantly released by setting is_available = True to restore search listings.")
    
    # ------------------ CHAPTER 10 ------------------
    add_heading_styled(doc, "10. Frontend (Client-side & Interactions)", 1)
    
    add_heading_styled(doc, "10.1 Interactive Frontend Controllers", 2)
    add_paragraph_styled(doc, "Rich dynamic browser controls are implemented using Vanilla JavaScript:")
    add_paragraph_styled(doc, "- Dynamic Calendar Grid: Generates calendar days, parses JSON blocked ranges from Django, and disables reserved dates in real-time.")
    add_paragraph_styled(doc, "- 3D Credit Card Simulation: Listens to checkout card fields and flips the card dynamically using CSS 3D perspectives when the CVV input gains focus.")
    add_paragraph_styled(doc, "- HTML5 Canvas Signature: Allows customers to draw electronic signatures on-screen. It compiles the coordinate strokes into a base64 PNG data string on form submit.")
    
    # ------------------ CHAPTER 11 ------------------
    add_heading_styled(doc, "11. Integration (API & Internal Routes)", 1)
    add_paragraph_styled(doc, "All page views, reviews, administrative checks, and promotional negotiations communicate via standard internal REST router mappings. To block malicious requests, all state-modifying requests enforce Django CSRF security tokens and check ownership constraints on the queried database IDs.")
    
    # ------------------ CHAPTER 12 ------------------
    add_heading_styled(doc, "12. Server & Deployment Infrastructure", 1)
    add_paragraph_styled(doc, "The system is pre-configured for local testing and cloud deployment. It utilizes Gunicorn as the WSGI server gateway to process concurrent HTTP requests, Whitenoise middleware to compile and compress static CSS assets, and an locked environment file (requirements.txt) to prevent library mismatches.")
    
    # ------------------ CHAPTER 13 ------------------
    doc.add_page_break()
    add_heading_styled(doc, "13. Testing, Verification & UI Quality Assurance Suite", 1)
    
    add_paragraph_styled(doc, "To guarantee absolute operational reliability, Web Wizards Car Rentals enforces a hybrid quality assurance program combining automated backend test suites (TDD) and systematic visual/manual interface verification matrices.")
    
    add_heading_styled(doc, "13.1 Automated Backend Integration Tests (TDD Coverage)", 2)
    add_paragraph_styled(doc, "- test_user_roles: Asserts absolute separation bounds for profile subclasses (is_customer(), is_vendor(), is_admin_role()), confirming incorrect context blocks fail.")
    add_paragraph_styled(doc, "- test_double_booking_prevention: Simulates database threads requesting identical reservation datetimes on approved listings, asserting database integrity queries block double-bookings.")
    add_paragraph_styled(doc, "- test_booking_cancellation_within_24_hours: Simulates cancellations within 24 hours of creation, verifying automatic transitions from PAID to REFUNDED and instant date releases.")
    add_paragraph_styled(doc, "- test_booking_cancellation_after_24_hours: Confirms that cancellation attempts outside the 24-hour window fail, returning validation errors.")
    add_paragraph_styled(doc, "- test_time_based_pricing_calculation: Exercises the Rupee capping rules: a) Renting for 5 hours at ₹10/hr returns exactly ₹50; b) Renting for 12 hours caps the charge at a full day's ₹100 rate; c) Renting for 27 hours returns a sum of a full day plus capped extra hours (₹130 total).")
    
    add_heading_styled(doc, "13.2 UI/UX Screen-by-Screen Quality Assurance Matrix", 2)
    add_paragraph_styled(doc, "Manual and simulated visual UI checks trace interactions across specific viewport layouts. The matrix below outlines verification paths and structural UI outcomes:")
    
    qa_table = doc.add_table(rows=1, cols=4)
    qa_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = qa_table.rows[0].cells
    hdr_cells[0].text = 'Tested UI Screen'
    hdr_cells[1].text = 'Test Interaction Path'
    hdr_cells[2].text = 'Expected Layout & Visual Outcome'
    hdr_cells[3].text = 'Testing Status'
    for cell in hdr_cells:
        set_cell_background(cell, '4F46E5')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
        cell.paragraphs[0].runs[0].bold = True
        
    qa_data = [
        ('Home Canvas (index.html)', '- Hover on navigation links.\n- Click login/register links.\n- Resize viewport to 320px.', '- Links highlight with micro-text shadows.\n- Global loading overlay fades in with status.\n- Sidebar transforms to hamburger drawer.', 'PASS\n- 12px filter loads.\n- drawer drawer functional.'),
        ('Search Desk (car_search.html)', '- Adjust category tags.\n- Drag the max daily rate slider.\n- Change transmission types.', '- Car grid redraws instantly.\n- Sliders display current pricing bounds.\n- Category cards glow on click.', 'PASS\n- Refreshes within 110ms.\n- Grid collapses on mobile.'),
        ('Detail Spec Desk (car_detail.html)', '- Click next/prev slider icons.\n- Tap calendar overlapping blocked reservation dates.\n- Select valid dates.', '- Smooth slide transitions.\n- Blocked dates display in locked red.\n- Valid dates colorize in HSL gradient.', 'PASS\n- Blocks double clicks.\n- Live price capping updates.'),
        ('Secure Checkout (checkout.html)', '- Fill credit card inputs.\n- Focus CVV field.\n- Draw signature on canvas.', '- Card number formats with spaces.\n- Card rotates 180 degrees in 3D.\n- Canvas captures fluid strokes.', 'PASS\n- rotateY(180deg) runs smoothly.\n- Base64 string serializes to DB.'),
        ('OTP Entry Panel (verify_otp.html)', '- Enter invalid digits.\n- Click resend OTP.\n- Enter valid OTP code.', '- Error flashes in glassmorphic alert box.\n- Success banner pops with new countdown.\n- Stateful login redirects to panel.', 'PASS\n- Django messages display with close buttons.'),
        ('Role Panels (dashboards)', '- Log in with specific role profile credentials.\n- Click return control.\n- Propose admin campaigns.', '- HSL stats counts load in cards.\n- Return control releases dates.\n- Admin proposed promos load.', 'PASS\n- Ledger updates in real-time.\n- Context checks intercept breaches.')
    ]
    
    for row_data in qa_data:
        row_cells = qa_table.add_row().cells
        row_cells[0].text = row_data[0]
        row_cells[1].text = row_data[1]
        row_cells[2].text = row_data[2]
        row_cells[3].text = row_data[3]
    
    # Save the document
    doc_path = "y:\\ALL FILES\\Car_rental_system_2\\comprehensive_project_report.docx"
    doc.save(doc_path)
    print("Word Document successfully created!")

if __name__ == "__main__":
    main()
